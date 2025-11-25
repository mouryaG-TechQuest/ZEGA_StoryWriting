"""
ZEGA Document Parser Core
Convert documents (PDF, DOCX, TXT, etc.) to structured scene descriptions
"""

import os
import re
import json
import tempfile
import asyncio
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
from datetime import datetime
from abc import ABC, abstractmethod


@dataclass
class ParsedScene:
    """Represents a parsed scene from document"""
    scene_number: int
    title: str
    content: str
    characters: List[str] = field(default_factory=list)
    location: str = ""
    time_of_day: str = ""
    mood: str = ""
    summary: str = ""
    start_position: int = 0
    end_position: int = 0


@dataclass
class ParsedDocument:
    """Represents a fully parsed document"""
    success: bool
    filename: str = ""
    total_scenes: int = 0
    scenes: List[ParsedScene] = field(default_factory=list)
    characters_found: List[str] = field(default_factory=list)
    genre_hints: List[str] = field(default_factory=list)
    word_count: int = 0
    error: Optional[str] = None


class DocumentExtractor(ABC):
    """Base class for document extractors"""
    
    @abstractmethod
    def extract(self, file_path: str) -> str:
        """Extract text from document"""
        pass
    
    @abstractmethod
    def supports(self, extension: str) -> bool:
        """Check if this extractor supports the file type"""
        pass


class TextExtractor(DocumentExtractor):
    """Extract text from plain text files"""
    
    def supports(self, extension: str) -> bool:
        return extension.lower() in ['.txt', '.text', '.md', '.markdown']
    
    def extract(self, file_path: str) -> str:
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        raise ValueError("Could not decode text file with any supported encoding")


class PDFExtractor(DocumentExtractor):
    """Extract text from PDF files"""
    
    def supports(self, extension: str) -> bool:
        return extension.lower() == '.pdf'
    
    def extract(self, file_path: str) -> str:
        try:
            import PyPDF2
            text_parts = []
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text_parts.append(page.extract_text())
            return '\n\n'.join(text_parts)
        except ImportError:
            # Fallback to pdfplumber
            try:
                import pdfplumber
                text_parts = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text_parts.append(page.extract_text() or '')
                return '\n\n'.join(text_parts)
            except ImportError:
                raise ImportError("Install PyPDF2 or pdfplumber: pip install PyPDF2 pdfplumber")


class DocxExtractor(DocumentExtractor):
    """Extract text from DOCX files"""
    
    def supports(self, extension: str) -> bool:
        return extension.lower() in ['.docx', '.doc']
    
    def extract(self, file_path: str) -> str:
        try:
            from docx import Document
            doc = Document(file_path)
            text_parts = []
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            return '\n'.join(text_parts)
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")


class RTFExtractor(DocumentExtractor):
    """Extract text from RTF files"""
    
    def supports(self, extension: str) -> bool:
        return extension.lower() == '.rtf'
    
    def extract(self, file_path: str) -> str:
        try:
            from striprtf.striprtf import rtf_to_text
            with open(file_path, 'r') as f:
                rtf_content = f.read()
            return rtf_to_text(rtf_content)
        except ImportError:
            raise ImportError("Install striprtf: pip install striprtf")


class EPUBExtractor(DocumentExtractor):
    """Extract text from EPUB files"""
    
    def supports(self, extension: str) -> bool:
        return extension.lower() == '.epub'
    
    def extract(self, file_path: str) -> str:
        try:
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup
            
            book = epub.read_epub(file_path)
            text_parts = []
            
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    text_parts.append(soup.get_text())
            
            return '\n\n'.join(text_parts)
        except ImportError:
            raise ImportError("Install ebooklib and beautifulsoup4: pip install ebooklib beautifulsoup4")


class HTMLExtractor(DocumentExtractor):
    """Extract text from HTML files"""
    
    def supports(self, extension: str) -> bool:
        return extension.lower() in ['.html', '.htm']
    
    def extract(self, file_path: str) -> str:
        try:
            from bs4 import BeautifulSoup
            with open(file_path, 'r', encoding='utf-8') as f:
                soup = BeautifulSoup(f.read(), 'html.parser')
            # Remove scripts and styles
            for element in soup(['script', 'style']):
                element.decompose()
            return soup.get_text()
        except ImportError:
            raise ImportError("Install beautifulsoup4: pip install beautifulsoup4")


class SceneParser:
    """Parse extracted text into structured scenes"""
    
    # Common scene break patterns
    SCENE_BREAK_PATTERNS = [
        r'\n\s*CHAPTER\s+(\d+|[IVXLCDM]+)[:\s]*(.+?)?\n',  # Chapter headers
        r'\n\s*SCENE\s+(\d+)[:\s]*(.+?)?\n',  # Scene headers
        r'\n\s*ACT\s+(\d+|[IVXLCDM]+)[:\s]*(.+?)?\n',  # Act headers
        r'\n\s*PART\s+(\d+|[IVXLCDM]+)[:\s]*(.+?)?\n',  # Part headers
        r'\n\s*#{1,3}\s+(.+?)\n',  # Markdown headers
        r'\n\s*\*\s*\*\s*\*\s*\n',  # Asterisk breaks
        r'\n\s*~\s*~\s*~\s*\n',  # Tilde breaks
        r'\n\s*-{3,}\s*\n',  # Dash breaks
        r'\n\s*_{3,}\s*\n',  # Underscore breaks
        r'\n{3,}',  # Multiple blank lines
    ]
    
    # Character name patterns (dialogue attribution)
    CHARACTER_PATTERNS = [
        r'"[^"]+"\s*said\s+(\w+)',  # "text" said John
        r'(\w+)\s+said[,.]',  # John said
        r'"[^"]+"\s+(\w+)\s+replied',  # "text" John replied
        r'(\w+)\s+asked[,.]',  # John asked
        r'(\w+)\s+whispered[,.]',  # John whispered
        r'(\w+)\s+shouted[,.]',  # John shouted
        r'(\w+):\s*"[^"]+"',  # John: "text" (screenplay format)
        r'^([A-Z][A-Z\s]+)$',  # ALL CAPS names (screenplay)
    ]
    
    # Location patterns
    LOCATION_PATTERNS = [
        r'INT\.\s*(.+?)\s*[-–]',  # INT. LOCATION -
        r'EXT\.\s*(.+?)\s*[-–]',  # EXT. LOCATION -
        r'at\s+the\s+(.+?)[,.]',  # at the [location]
        r'in\s+the\s+(.+?)[,.]',  # in the [location]
        r'inside\s+(?:the\s+)?(.+?)[,.]',  # inside [the location]
    ]
    
    # Time of day patterns
    TIME_PATTERNS = [
        r'\b(morning|dawn|sunrise)\b',
        r'\b(afternoon|midday|noon)\b',
        r'\b(evening|dusk|sunset)\b',
        r'\b(night|midnight)\b',
        r'[-–]\s*(DAY|NIGHT|CONTINUOUS|LATER)',  # Screenplay format
    ]
    
    # Mood/atmosphere keywords
    MOOD_KEYWORDS = {
        'tense': ['tension', 'anxious', 'nervous', 'fear', 'dread', 'worried'],
        'romantic': ['love', 'romantic', 'passionate', 'tender', 'intimate'],
        'action': ['chase', 'fight', 'run', 'explosion', 'battle', 'combat'],
        'mystery': ['mysterious', 'enigma', 'secret', 'hidden', 'clue'],
        'comedy': ['laugh', 'funny', 'joke', 'amusing', 'hilarious'],
        'dark': ['dark', 'gloomy', 'shadow', 'sinister', 'ominous'],
        'peaceful': ['calm', 'serene', 'quiet', 'peaceful', 'tranquil'],
        'sad': ['sad', 'grief', 'sorrow', 'tears', 'mourning', 'loss']
    }
    
    def __init__(self):
        self.extractors = [
            TextExtractor(),
            PDFExtractor(),
            DocxExtractor(),
            RTFExtractor(),
            EPUBExtractor(),
            HTMLExtractor()
        ]
    
    def get_extractor(self, file_path: str) -> Optional[DocumentExtractor]:
        """Get appropriate extractor for file type"""
        ext = Path(file_path).suffix
        for extractor in self.extractors:
            if extractor.supports(ext):
                return extractor
        return None
    
    def _find_scene_breaks(self, text: str) -> List[Tuple[int, str]]:
        """Find scene break positions and their titles"""
        breaks = []
        
        for pattern in self.SCENE_BREAK_PATTERNS:
            for match in re.finditer(pattern, text, re.MULTILINE | re.IGNORECASE):
                pos = match.start()
                # Try to extract title
                title = ""
                groups = match.groups()
                if groups:
                    title = ' '.join(str(g) for g in groups if g).strip()
                breaks.append((pos, title))
        
        # Sort and deduplicate (keep unique positions)
        breaks.sort(key=lambda x: x[0])
        unique_breaks = []
        last_pos = -100
        for pos, title in breaks:
            if pos - last_pos > 50:  # Minimum gap between scenes
                unique_breaks.append((pos, title))
                last_pos = pos
        
        return unique_breaks
    
    def _extract_characters(self, text: str) -> List[str]:
        """Extract character names from text"""
        characters = set()
        
        for pattern in self.CHARACTER_PATTERNS:
            for match in re.finditer(pattern, text, re.MULTILINE):
                name = match.group(1).strip()
                # Filter out common non-name words
                if len(name) > 1 and name.lower() not in ['the', 'a', 'an', 'he', 'she', 'it', 'they']:
                    characters.add(name.title())
        
        return list(characters)[:20]  # Limit to 20 characters
    
    def _extract_location(self, text: str) -> str:
        """Extract primary location from scene text"""
        for pattern in self.LOCATION_PATTERNS:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        return ""
    
    def _extract_time(self, text: str) -> str:
        """Extract time of day from scene text"""
        for pattern in self.TIME_PATTERNS:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1).lower()
        return ""
    
    def _detect_mood(self, text: str) -> str:
        """Detect scene mood from keywords"""
        text_lower = text.lower()
        mood_scores = {}
        
        for mood, keywords in self.MOOD_KEYWORDS.items():
            score = sum(1 for kw in keywords if kw in text_lower)
            if score > 0:
                mood_scores[mood] = score
        
        if mood_scores:
            return max(mood_scores, key=mood_scores.get)
        return ""
    
    def _generate_summary(self, text: str, max_words: int = 50) -> str:
        """Generate a brief summary of scene content"""
        # Get first few sentences
        sentences = re.split(r'[.!?]', text)
        summary_parts = []
        word_count = 0
        
        for sentence in sentences:
            words = sentence.split()
            if word_count + len(words) <= max_words:
                summary_parts.append(sentence.strip())
                word_count += len(words)
            else:
                break
        
        summary = '. '.join(summary_parts)
        if summary and not summary.endswith('.'):
            summary += '.'
        
        return summary
    
    def parse_text(self, text: str, filename: str = "document") -> ParsedDocument:
        """Parse raw text into structured scenes"""
        
        if not text or not text.strip():
            return ParsedDocument(
                success=False,
                filename=filename,
                error="Empty or invalid document content"
            )
        
        # Clean text
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Find scene breaks
        breaks = self._find_scene_breaks(text)
        
        # If no breaks found, treat as single scene
        if not breaks:
            breaks = [(0, "Scene 1")]
        
        # Add end position
        breaks.append((len(text), ""))
        
        scenes = []
        all_characters = set()
        
        for i in range(len(breaks) - 1):
            start_pos, title = breaks[i]
            end_pos = breaks[i + 1][0]
            
            scene_text = text[start_pos:end_pos].strip()
            
            # Skip very short segments
            if len(scene_text) < 50:
                continue
            
            # Generate title if not found
            if not title:
                title = f"Scene {len(scenes) + 1}"
            
            # Extract metadata
            characters = self._extract_characters(scene_text)
            all_characters.update(characters)
            
            scene = ParsedScene(
                scene_number=len(scenes) + 1,
                title=title,
                content=scene_text,
                characters=characters,
                location=self._extract_location(scene_text),
                time_of_day=self._extract_time(scene_text),
                mood=self._detect_mood(scene_text),
                summary=self._generate_summary(scene_text),
                start_position=start_pos,
                end_position=end_pos
            )
            scenes.append(scene)
        
        # Detect genre hints
        genre_hints = self._detect_genre_hints(text)
        
        return ParsedDocument(
            success=True,
            filename=filename,
            total_scenes=len(scenes),
            scenes=scenes,
            characters_found=list(all_characters),
            genre_hints=genre_hints,
            word_count=len(text.split())
        )
    
    def _detect_genre_hints(self, text: str) -> List[str]:
        """Detect possible genres from text content"""
        genre_keywords = {
            'fantasy': ['magic', 'wizard', 'dragon', 'sword', 'kingdom', 'elf', 'dwarf', 'spell'],
            'sci-fi': ['spaceship', 'alien', 'robot', 'laser', 'planet', 'galaxy', 'android'],
            'romance': ['love', 'heart', 'kiss', 'romantic', 'passion', 'beloved'],
            'mystery': ['detective', 'clue', 'murder', 'suspect', 'investigate', 'crime'],
            'thriller': ['chase', 'danger', 'escape', 'agent', 'terrorist', 'conspiracy'],
            'horror': ['ghost', 'demon', 'haunted', 'blood', 'scream', 'monster', 'nightmare'],
            'drama': ['family', 'conflict', 'struggle', 'relationship', 'crisis'],
            'comedy': ['laugh', 'joke', 'funny', 'hilarious', 'comedy'],
            'action': ['fight', 'battle', 'explosion', 'weapon', 'attack', 'combat'],
            'historical': ['ancient', 'medieval', 'war', 'empire', 'century', 'kingdom']
        }
        
        text_lower = text.lower()
        genre_scores = {}
        
        for genre, keywords in genre_keywords.items():
            score = sum(text_lower.count(kw) for kw in keywords)
            if score > 2:  # Minimum threshold
                genre_scores[genre] = score
        
        # Return top 3 genres
        sorted_genres = sorted(genre_scores.keys(), key=lambda g: genre_scores[g], reverse=True)
        return sorted_genres[:3]


class ZegaDocParser:
    """
    Main document parser class
    Handles document extraction, parsing, and scene generation
    """
    
    def __init__(self, training_data_path: Optional[str] = None):
        self.parser = SceneParser()
        self.training_data_path = Path(training_data_path) if training_data_path else None
        
        if self.training_data_path:
            self.training_data_path.mkdir(parents=True, exist_ok=True)
    
    async def parse_file(self, file_path: str) -> ParsedDocument:
        """
        Parse a document file into structured scenes
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ParsedDocument with extracted scenes
        """
        path = Path(file_path)
        
        if not path.exists():
            return ParsedDocument(
                success=False,
                filename=path.name,
                error=f"File not found: {file_path}"
            )
        
        # Get appropriate extractor
        extractor = self.parser.get_extractor(file_path)
        
        if not extractor:
            return ParsedDocument(
                success=False,
                filename=path.name,
                error=f"Unsupported file format: {path.suffix}"
            )
        
        try:
            # Extract text
            text = extractor.extract(file_path)
            
            # Parse into scenes
            result = self.parser.parse_text(text, path.name)
            
            # Save for training
            if self.training_data_path and result.success:
                await self._save_training_data(result)
            
            return result
            
        except Exception as e:
            return ParsedDocument(
                success=False,
                filename=path.name,
                error=str(e)
            )
    
    async def parse_bytes(
        self,
        data: bytes,
        filename: str,
        content_type: Optional[str] = None
    ) -> ParsedDocument:
        """
        Parse document from bytes
        
        Args:
            data: Raw file bytes
            filename: Original filename (used for extension detection)
            content_type: Optional MIME type
            
        Returns:
            ParsedDocument with extracted scenes
        """
        # Write to temp file
        ext = Path(filename).suffix or self._guess_extension(content_type)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        
        try:
            return await self.parse_file(tmp_path)
        finally:
            # Clean up temp file
            try:
                os.unlink(tmp_path)
            except:
                pass
    
    async def parse_text_direct(self, text: str, title: str = "Untitled") -> ParsedDocument:
        """
        Parse raw text directly (no file)
        
        Args:
            text: Raw text content
            title: Document title
            
        Returns:
            ParsedDocument with extracted scenes
        """
        return self.parser.parse_text(text, title)
    
    def _guess_extension(self, content_type: Optional[str]) -> str:
        """Guess file extension from content type"""
        if not content_type:
            return '.txt'
        
        mime_map = {
            'application/pdf': '.pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
            'application/msword': '.doc',
            'application/rtf': '.rtf',
            'application/epub+zip': '.epub',
            'text/html': '.html',
            'text/plain': '.txt',
            'text/markdown': '.md'
        }
        
        return mime_map.get(content_type, '.txt')
    
    async def _save_training_data(self, result: ParsedDocument):
        """Save parsed document for training data"""
        if not self.training_data_path:
            return
        
        training_entry = {
            'filename': result.filename,
            'total_scenes': result.total_scenes,
            'word_count': result.word_count,
            'characters_found': result.characters_found,
            'genre_hints': result.genre_hints,
            'scenes': [
                {
                    'title': s.title,
                    'content_length': len(s.content),
                    'characters': s.characters,
                    'location': s.location,
                    'time_of_day': s.time_of_day,
                    'mood': s.mood,
                    'summary': s.summary
                }
                for s in result.scenes
            ],
            'parsed_at': datetime.now().isoformat()
        }
        
        # Save to JSONL file
        training_file = self.training_data_path / "parsing_examples.jsonl"
        with open(training_file, 'a', encoding='utf-8') as f:
            f.write(json.dumps(training_entry) + '\n')
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return ['.txt', '.md', '.pdf', '.docx', '.doc', '.rtf', '.epub', '.html', '.htm']
    
    async def enhance_scene_with_ai(
        self,
        scene: ParsedScene,
        style: str = "narrative"
    ) -> ParsedScene:
        """
        Enhance a parsed scene with AI-generated details
        
        This is a placeholder for AI enhancement - can be connected
        to the ZEGA main service for full scene enhancement.
        """
        # For now, just return the scene
        # In production, this would call the AI service
        return scene
    
    def scenes_to_story_format(self, result: ParsedDocument) -> List[Dict[str, Any]]:
        """
        Convert parsed document to story service format
        
        Returns list ready for story-service API
        """
        story_scenes = []
        
        for scene in result.scenes:
            story_scene = {
                'sceneNumber': scene.scene_number,
                'title': scene.title,
                'description': scene.content,
                'summary': scene.summary,
                'characters': [{'name': c} for c in scene.characters],
                'location': scene.location,
                'timeOfDay': scene.time_of_day,
                'mood': scene.mood,
                'metadata': {
                    'source': result.filename,
                    'parsed_at': datetime.now().isoformat()
                }
            }
            story_scenes.append(story_scene)
        
        return story_scenes
